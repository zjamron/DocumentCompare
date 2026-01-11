"""
Test script that demonstrates the document comparison logic.
This is a Python implementation of the same algorithm used in the C# library.
"""
import difflib
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX
import os

class DiffType:
    UNCHANGED = 'unchanged'
    INSERTED = 'inserted'
    DELETED = 'deleted'

class DiffSegment:
    def __init__(self, text, diff_type):
        self.text = text
        self.type = diff_type

def tokenize_to_words(text):
    """Split text into words."""
    return re.findall(r'\S+', text)

def get_paragraph_text(para):
    """Get plain text from a paragraph."""
    return para.text.strip()

def diff_paragraphs(original_text, modified_text):
    """Perform word-level diff between two paragraphs."""
    if not original_text and not modified_text:
        return []

    if not original_text:
        return [DiffSegment(modified_text, DiffType.INSERTED)]

    if not modified_text:
        return [DiffSegment(original_text, DiffType.DELETED)]

    # Tokenize to words
    original_words = tokenize_to_words(original_text)
    modified_words = tokenize_to_words(modified_text)

    # Use difflib for word-level diff
    matcher = difflib.SequenceMatcher(None, original_words, modified_words)
    segments = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            text = ' '.join(original_words[i1:i2])
            segments.append(DiffSegment(text, DiffType.UNCHANGED))
        elif tag == 'delete':
            text = ' '.join(original_words[i1:i2])
            segments.append(DiffSegment(text, DiffType.DELETED))
        elif tag == 'insert':
            text = ' '.join(modified_words[j1:j2])
            segments.append(DiffSegment(text, DiffType.INSERTED))
        elif tag == 'replace':
            # Delete old, insert new
            del_text = ' '.join(original_words[i1:i2])
            ins_text = ' '.join(modified_words[j1:j2])
            segments.append(DiffSegment(del_text, DiffType.DELETED))
            segments.append(DiffSegment(ins_text, DiffType.INSERTED))

    return segments

def calculate_similarity(text1, text2):
    """Calculate similarity between two texts (0-1)."""
    if text1 == text2:
        return 1.0
    if not text1 or not text2:
        return 0.0

    words1 = set(tokenize_to_words(text1.lower()))
    words2 = set(tokenize_to_words(text2.lower()))

    if not words1 and not words2:
        return 1.0
    if not words1 or not words2:
        return 0.0

    intersection = len(words1 & words2)
    union = len(words1 | words2)

    return intersection / union if union > 0 else 0.0

def align_paragraphs(original_paras, modified_paras):
    """Align paragraphs using LCS algorithm."""
    m, n = len(original_paras), len(modified_paras)

    # Build LCS table
    lcs = [[0] * (n + 1) for _ in range(m + 1)]

    for i in range(1, m + 1):
        for j in range(1, n + 1):
            orig_text = get_paragraph_text(original_paras[i-1])
            mod_text = get_paragraph_text(modified_paras[j-1])

            if calculate_similarity(orig_text, mod_text) >= 0.5:
                lcs[i][j] = lcs[i-1][j-1] + 1
            else:
                lcs[i][j] = max(lcs[i-1][j], lcs[i][j-1])

    # Backtrack to find alignment
    alignments = []
    i, j = m, n

    while i > 0 or j > 0:
        if i > 0 and j > 0:
            orig_text = get_paragraph_text(original_paras[i-1])
            mod_text = get_paragraph_text(modified_paras[j-1])
            if calculate_similarity(orig_text, mod_text) >= 0.5:
                alignments.append((i-1, j-1, DiffType.UNCHANGED))
                i -= 1
                j -= 1
                continue

        if j > 0 and (i == 0 or lcs[i][j-1] >= lcs[i-1][j]):
            alignments.append((-1, j-1, DiffType.INSERTED))
            j -= 1
        else:
            alignments.append((i-1, -1, DiffType.DELETED))
            i -= 1

    alignments.reverse()
    return alignments

def set_run_formatting(run, diff_type):
    """Apply redline formatting to a run."""
    if diff_type == DiffType.DELETED:
        run.font.strike = True
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red
    elif diff_type == DiffType.INSERTED:
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 255)  # Blue

def create_redlined_document(original_path, modified_path, output_path):
    """Create a redlined document showing differences."""
    original_doc = Document(original_path)
    modified_doc = Document(modified_path)

    # Create output document
    redlined_doc = Document()

    # Get paragraphs (excluding empty ones for alignment purposes, but we'll handle them)
    original_paras = list(original_doc.paragraphs)
    modified_paras = list(modified_doc.paragraphs)

    # Align paragraphs
    alignments = align_paragraphs(original_paras, modified_paras)

    # Statistics
    stats = {
        'insertions': 0,
        'deletions': 0,
        'unchanged': 0
    }

    # Generate redlined content
    for orig_idx, mod_idx, alignment_type in alignments:
        if alignment_type == DiffType.DELETED:
            # Entire paragraph deleted
            para = redlined_doc.add_paragraph()
            orig_para = original_paras[orig_idx]

            # Copy style if it's a heading
            if orig_para.style.name.startswith('Heading'):
                para.style = redlined_doc.styles[orig_para.style.name]

            text = get_paragraph_text(orig_para)
            if text:
                run = para.add_run(text)
                set_run_formatting(run, DiffType.DELETED)
                stats['deletions'] += len(tokenize_to_words(text))

        elif alignment_type == DiffType.INSERTED:
            # Entire paragraph inserted
            para = redlined_doc.add_paragraph()
            mod_para = modified_paras[mod_idx]

            # Copy style if it's a heading
            if mod_para.style.name.startswith('Heading'):
                para.style = redlined_doc.styles[mod_para.style.name]

            text = get_paragraph_text(mod_para)
            if text:
                run = para.add_run(text)
                set_run_formatting(run, DiffType.INSERTED)
                stats['insertions'] += len(tokenize_to_words(text))

        else:
            # Paragraph exists in both - do word-level diff
            orig_para = original_paras[orig_idx]
            mod_para = modified_paras[mod_idx]

            para = redlined_doc.add_paragraph()

            # Copy style if it's a heading
            if mod_para.style.name.startswith('Heading'):
                para.style = redlined_doc.styles[mod_para.style.name]

            orig_text = get_paragraph_text(orig_para)
            mod_text = get_paragraph_text(mod_para)

            segments = diff_paragraphs(orig_text, mod_text)

            for i, segment in enumerate(segments):
                if segment.text:
                    # Add space between segments (except at start)
                    if i > 0:
                        para.add_run(' ')

                    run = para.add_run(segment.text)
                    set_run_formatting(run, segment.type)

                    word_count = len(tokenize_to_words(segment.text))
                    if segment.type == DiffType.INSERTED:
                        stats['insertions'] += word_count
                    elif segment.type == DiffType.DELETED:
                        stats['deletions'] += word_count
                    else:
                        stats['unchanged'] += word_count

    # Save redlined document
    redlined_doc.save(output_path)

    return stats

def print_diff_summary(original_path, modified_path):
    """Print a text-based diff summary."""
    original_doc = Document(original_path)
    modified_doc = Document(modified_path)

    original_paras = list(original_doc.paragraphs)
    modified_paras = list(modified_doc.paragraphs)

    alignments = align_paragraphs(original_paras, modified_paras)

    print("\n" + "="*80)
    print("DOCUMENT COMPARISON SUMMARY")
    print("="*80)

    para_num = 0
    for orig_idx, mod_idx, alignment_type in alignments:
        para_num += 1

        if alignment_type == DiffType.DELETED:
            text = get_paragraph_text(original_paras[orig_idx])
            if text:
                print(f"\n[DELETED] Paragraph {para_num}:")
                print(f"  - {text[:100]}{'...' if len(text) > 100 else ''}")

        elif alignment_type == DiffType.INSERTED:
            text = get_paragraph_text(modified_paras[mod_idx])
            if text:
                print(f"\n[INSERTED] Paragraph {para_num}:")
                print(f"  + {text[:100]}{'...' if len(text) > 100 else ''}")

        else:
            orig_text = get_paragraph_text(original_paras[orig_idx])
            mod_text = get_paragraph_text(modified_paras[mod_idx])

            if orig_text != mod_text and (orig_text or mod_text):
                segments = diff_paragraphs(orig_text, mod_text)
                has_changes = any(s.type != DiffType.UNCHANGED for s in segments)

                if has_changes:
                    print(f"\n[MODIFIED] Paragraph {para_num}:")
                    for segment in segments:
                        if segment.type == DiffType.DELETED:
                            print(f"  - {segment.text}")
                        elif segment.type == DiffType.INSERTED:
                            print(f"  + {segment.text}")

if __name__ == '__main__':
    samples_dir = os.path.dirname(os.path.abspath(__file__))

    original_path = os.path.join(samples_dir, 'contract_v1.docx')
    modified_path = os.path.join(samples_dir, 'contract_v2.docx')
    output_path = os.path.join(samples_dir, 'contract_redline.docx')

    print("Document Comparison Test")
    print("="*50)
    print(f"Original: {original_path}")
    print(f"Modified: {modified_path}")
    print(f"Output:   {output_path}")

    # Print text-based diff summary
    print_diff_summary(original_path, modified_path)

    # Create redlined document
    print("\n" + "="*80)
    print("GENERATING REDLINED DOCUMENT")
    print("="*80)

    stats = create_redlined_document(original_path, modified_path, output_path)

    print(f"\nRedlined document created: {output_path}")
    print(f"\nStatistics:")
    print(f"  Words inserted (blue bold):      {stats['insertions']}")
    print(f"  Words deleted (red strikethrough): {stats['deletions']}")
    print(f"  Words unchanged:                   {stats['unchanged']}")

    total = stats['insertions'] + stats['deletions'] + stats['unchanged']
    if total > 0:
        change_pct = (stats['insertions'] + stats['deletions']) * 100 / total
        print(f"  Change percentage:                 {change_pct:.1f}%")

    print("\n" + "="*80)
    print("TEST COMPLETE")
    print("="*80)
    print("\nOpen the redlined document to see:")
    print("  - Deleted text in RED with strikethrough")
    print("  - Inserted text in BLUE and BOLD")
