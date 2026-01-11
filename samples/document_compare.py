"""
Universal Document Comparison Tool

Supports:
- Word (.docx) to Word comparison
- PDF to PDF comparison
- Word to PDF comparison (and vice versa)
- Output to Word or PDF format
- Move detection (green highlighting)
"""

import os
import sys
import difflib
import re
import shutil
from typing import List, Tuple, Optional, Dict
from dataclasses import dataclass

# Move detection settings
MOVE_SIMILARITY_THRESHOLD = 0.85
MIN_MOVE_WORDS = 3

# Import Word support
from docx import Document
from docx.shared import RGBColor

# Import PDF support
from pdf_support import (
    PdfParser, PdfGenerator, PdfDocumentAdapter,
    ExtractedParagraph, extract_paragraphs_from_pdf
)


@dataclass
class ComparisonResult:
    """Result of a document comparison."""
    output_path: str
    insertions: int
    deletions: int
    moves: int
    unchanged: int
    success: bool
    error: Optional[str] = None


def get_file_type(filepath: str) -> str:
    """Determine file type from extension."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.docx':
        return 'word'
    elif ext == '.pdf':
        return 'pdf'
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def tokenize(text: str) -> List[str]:
    """Split text into words while preserving whitespace."""
    return re.findall(r'\S+|\s+', text)


def diff_texts(original_text: str, modified_text: str) -> List[Tuple[str, str]]:
    """
    Compute word-level diff between two texts.
    Returns list of (text, type) where type is 'equal', 'delete', or 'insert'
    """
    orig_words = tokenize(original_text)
    mod_words = tokenize(modified_text)

    matcher = difflib.SequenceMatcher(None, orig_words, mod_words, autojunk=False)
    result = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            text = ''.join(orig_words[i1:i2])
            if text:
                result.append((text, 'equal'))
        elif tag == 'delete':
            text = ''.join(orig_words[i1:i2])
            if text:
                result.append((text, 'delete'))
        elif tag == 'insert':
            text = ''.join(mod_words[j1:j2])
            if text:
                result.append((text, 'insert'))
        elif tag == 'replace':
            del_text = ''.join(orig_words[i1:i2])
            ins_text = ''.join(mod_words[j1:j2])
            if del_text:
                result.append((del_text, 'delete'))
            if ins_text:
                result.append((ins_text, 'insert'))

    return result


def calculate_similarity(text1: str, text2: str) -> float:
    """Calculate similarity between two texts using multiple methods."""
    if text1 == text2:
        return 1.0
    if not text1 or not text2:
        return 0.0

    text1 = text1.strip()
    text2 = text2.strip()

    if not text1 and not text2:
        return 1.0
    if not text1 or not text2:
        return 0.0

    # Word-based Jaccard similarity
    words1 = set(text1.lower().split())
    words2 = set(text2.lower().split())

    word_sim = 0.0
    if words1 and words2:
        intersection = len(words1 & words2)
        union = len(words1 | words2)
        word_sim = intersection / union if union > 0 else 0.0

    # Character sequence similarity
    seq_sim = difflib.SequenceMatcher(None, text1.lower(), text2.lower()).ratio()

    return max(word_sim, seq_sim)


def normalize_text_for_move(text: str) -> str:
    """Normalize text for move detection comparison."""
    text = text.lower().strip()
    text = re.sub(r'\s+', ' ', text)
    return text


def detect_word_level_moves(diff_segments: List[Tuple[str, str]]) -> List[Tuple[str, str]]:
    """
    Detect moves within word-level diff segments.
    Converts delete/insert pairs to move_source/move_dest when text is similar.
    """
    # Collect deletions and insertions
    deletions = []  # (index, text, normalized)
    insertions = []  # (index, text, normalized)

    for i, (text, seg_type) in enumerate(diff_segments):
        words = len(text.split())
        if seg_type == 'delete' and words >= MIN_MOVE_WORDS:
            deletions.append((i, text, normalize_text_for_move(text)))
        elif seg_type == 'insert' and words >= MIN_MOVE_WORDS:
            insertions.append((i, text, normalize_text_for_move(text)))

    if not deletions or not insertions:
        return diff_segments

    # Find matching moves
    moves = {}  # deletion_idx -> insertion_idx
    used_insertions = set()

    # Sort deletions by word count descending
    sorted_deletions = sorted(deletions, key=lambda x: len(x[1].split()), reverse=True)

    for del_idx, del_text, del_norm in sorted_deletions:
        best_match = None
        best_similarity = 0

        for ins_idx, ins_text, ins_norm in insertions:
            if ins_idx in used_insertions:
                continue

            similarity = calculate_similarity(del_norm, ins_norm)

            if similarity >= MOVE_SIMILARITY_THRESHOLD and similarity > best_similarity:
                best_similarity = similarity
                best_match = ins_idx

        if best_match is not None:
            moves[del_idx] = best_match
            used_insertions.add(best_match)

    if not moves:
        return diff_segments

    # Create new segments with move markers
    result = []
    for i, (text, seg_type) in enumerate(diff_segments):
        if i in moves:
            result.append((text, 'move_source'))
        elif i in used_insertions:
            result.append((text, 'move_dest'))
        else:
            result.append((text, seg_type))

    return result


def align_paragraphs(orig_texts: List[str], mod_texts: List[str]) -> List[Tuple[int, int, str]]:
    """Align paragraphs using LCS algorithm."""
    m, n = len(orig_texts), len(mod_texts)

    # Build LCS table
    lcs = [[0] * (n + 1) for _ in range(m + 1)]

    for i in range(1, m + 1):
        for j in range(1, n + 1):
            if calculate_similarity(orig_texts[i-1], mod_texts[j-1]) >= 0.4:
                lcs[i][j] = lcs[i-1][j-1] + 1
            else:
                lcs[i][j] = max(lcs[i-1][j], lcs[i][j-1])

    # Backtrack
    alignments = []
    i, j = m, n

    while i > 0 or j > 0:
        if i > 0 and j > 0:
            if calculate_similarity(orig_texts[i-1], mod_texts[j-1]) >= 0.4:
                alignments.append((i-1, j-1, 'match'))
                i -= 1
                j -= 1
                continue

        if j > 0 and (i == 0 or lcs[i][j-1] >= lcs[i-1][j]):
            alignments.append((-1, j-1, 'insert'))
            j -= 1
        else:
            alignments.append((i-1, -1, 'delete'))
            i -= 1

    alignments.reverse()
    return alignments


def extract_paragraphs_from_word(doc_path: str) -> List[Tuple[str, dict]]:
    """Extract paragraphs from Word document with metadata."""
    doc = Document(doc_path)
    result = []
    for para in doc.paragraphs:
        metadata = {
            'is_heading': para.style.name.startswith('Heading') if para.style else False
        }
        result.append((para.text, metadata))
    return result


def extract_tables_from_word(doc_path: str) -> List[List[List[str]]]:
    """
    Extract tables from Word document.
    Returns list of tables, each table is a list of rows, each row is a list of cell texts.
    """
    doc = Document(doc_path)
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = '\n'.join(p.text for p in cell.paragraphs)
                row_data.append(cell_text)
            table_data.append(row_data)
        tables.append(table_data)
    return tables


def extract_tables_from_document(doc_path: str) -> List[List[List[str]]]:
    """Extract tables from any supported document type."""
    file_type = get_file_type(doc_path)

    if file_type == 'word':
        return extract_tables_from_word(doc_path)
    elif file_type == 'pdf':
        # PDF table extraction is complex - for now return empty
        # Future: could use tabula-py or camelot
        return []


def compare_table_rows(orig_rows: List[List[str]], mod_rows: List[List[str]]) -> List[Tuple[int, int, str]]:
    """Align rows between two tables using LCS algorithm."""
    m, n = len(orig_rows), len(mod_rows)

    def get_row_text(row):
        return ' | '.join(row)

    # Build LCS table
    lcs = [[0] * (n + 1) for _ in range(m + 1)]

    for i in range(1, m + 1):
        for j in range(1, n + 1):
            orig_text = get_row_text(orig_rows[i-1])
            mod_text = get_row_text(mod_rows[j-1])

            if calculate_similarity(orig_text, mod_text) >= 0.4:
                lcs[i][j] = lcs[i-1][j-1] + 1
            else:
                lcs[i][j] = max(lcs[i-1][j], lcs[i][j-1])

    # Backtrack
    alignments = []
    i, j = m, n

    while i > 0 or j > 0:
        if i > 0 and j > 0:
            orig_text = get_row_text(orig_rows[i-1])
            mod_text = get_row_text(mod_rows[j-1])
            if calculate_similarity(orig_text, mod_text) >= 0.4:
                alignments.append((i-1, j-1, 'match'))
                i -= 1
                j -= 1
                continue

        if j > 0 and (i == 0 or lcs[i][j-1] >= lcs[i-1][j]):
            alignments.append((-1, j-1, 'insert'))
            j -= 1
        else:
            alignments.append((i-1, -1, 'delete'))
            i -= 1

    alignments.reverse()
    return alignments


def diff_table(orig_table: List[List[str]], mod_table: List[List[str]], stats: dict) -> List[dict]:
    """
    Compare two tables and return diff results for each row.
    """
    results = []
    row_alignments = compare_table_rows(orig_table, mod_table)

    for orig_row_idx, mod_row_idx, align_type in row_alignments:
        if align_type == 'match' and orig_row_idx >= 0 and mod_row_idx >= 0:
            orig_row = orig_table[orig_row_idx]
            mod_row = mod_table[mod_row_idx]

            # Compare cells
            row_segments = []
            max_cols = max(len(orig_row), len(mod_row))

            for col_idx in range(max_cols):
                if col_idx < len(orig_row) and col_idx < len(mod_row):
                    orig_cell = orig_row[col_idx]
                    mod_cell = mod_row[col_idx]

                    if orig_cell.strip() != mod_cell.strip():
                        # Cell changed - compute diff
                        cell_diff = diff_texts(orig_cell, mod_cell)
                        cell_diff = detect_word_level_moves(cell_diff)

                        for text, seg_type in cell_diff:
                            words = len(text.split())
                            if seg_type == 'insert':
                                stats['insertions'] += words
                            elif seg_type == 'delete':
                                stats['deletions'] += words
                            elif seg_type in ('move_source', 'move_dest'):
                                stats['moves'] += words
                            else:
                                stats['unchanged'] += words

                        row_segments.extend(cell_diff)
                    else:
                        row_segments.append((mod_cell, 'equal'))
                        stats['unchanged'] += len(mod_cell.split())

                    # Add cell separator
                    if col_idx < max_cols - 1:
                        row_segments.append((' | ', 'equal'))

                elif col_idx < len(mod_row):
                    # New column
                    row_segments.append((mod_row[col_idx], 'insert'))
                    stats['insertions'] += len(mod_row[col_idx].split())
                    if col_idx < max_cols - 1:
                        row_segments.append((' | ', 'equal'))

            results.append({'segments': row_segments, 'is_heading': False, 'is_table_row': True})

        elif align_type == 'insert' and mod_row_idx >= 0:
            # New row
            mod_row = mod_table[mod_row_idx]
            row_text = ' | '.join(mod_row)
            results.append({'segments': [(row_text, 'insert')], 'is_heading': False, 'is_table_row': True})
            stats['insertions'] += len(row_text.split())

        elif align_type == 'delete' and orig_row_idx >= 0:
            # Deleted row
            orig_row = orig_table[orig_row_idx]
            row_text = ' | '.join(orig_row)
            results.append({'segments': [(row_text, 'delete')], 'is_heading': False, 'is_table_row': True})
            stats['deletions'] += len(row_text.split())

    return results


def extract_paragraphs_from_document(doc_path: str) -> List[Tuple[str, dict]]:
    """Extract paragraphs from any supported document type."""
    file_type = get_file_type(doc_path)

    if file_type == 'word':
        return extract_paragraphs_from_word(doc_path)
    elif file_type == 'pdf':
        with PdfParser(doc_path) as parser:
            extracted = parser.extract_paragraphs()
            return [(p.text, {'is_heading': p.is_heading}) for p in extracted]


def compare_documents(
    original_path: str,
    modified_path: str,
    output_path: str,
    output_format: Optional[str] = None
) -> ComparisonResult:
    """
    Compare two documents and generate a redlined output.

    Args:
        original_path: Path to original document (Word or PDF)
        modified_path: Path to modified document (Word or PDF)
        output_path: Path for output redlined document
        output_format: 'word' or 'pdf' (auto-detected from output_path if not specified)

    Returns:
        ComparisonResult with statistics and status
    """
    try:
        # Determine output format
        if output_format is None:
            output_format = get_file_type(output_path)

        print(f"Original: {original_path} ({get_file_type(original_path)})")
        print(f"Modified: {modified_path} ({get_file_type(modified_path)})")
        print(f"Output: {output_path} ({output_format})")
        print()

        # Extract paragraphs from both documents
        print("Extracting paragraphs...")
        orig_paras = extract_paragraphs_from_document(original_path)
        mod_paras = extract_paragraphs_from_document(modified_path)

        print(f"  Original: {len(orig_paras)} paragraphs")
        print(f"  Modified: {len(mod_paras)} paragraphs")

        # Extract tables
        print("Extracting tables...")
        orig_tables = extract_tables_from_document(original_path)
        mod_tables = extract_tables_from_document(modified_path)

        print(f"  Original: {len(orig_tables)} tables")
        print(f"  Modified: {len(mod_tables)} tables")

        # Get just the text for alignment
        orig_texts = [text for text, _ in orig_paras]
        mod_texts = [text for text, _ in mod_paras]

        # Align paragraphs
        print("Aligning paragraphs...")
        alignments = align_paragraphs(orig_texts, mod_texts)

        # Compute diffs
        print("Computing differences...")
        diff_results = []
        stats = {'insertions': 0, 'deletions': 0, 'moves': 0, 'unchanged': 0}

        # First pass: collect all diffs
        temp_results = []
        for orig_idx, mod_idx, align_type in alignments:
            if align_type == 'match':
                orig_text = orig_texts[orig_idx]
                mod_text = mod_texts[mod_idx]
                mod_meta = mod_paras[mod_idx][1]

                if orig_text.strip() != mod_text.strip():
                    segments = diff_texts(orig_text, mod_text)
                    # Apply word-level move detection
                    segments = detect_word_level_moves(segments)
                else:
                    segments = [(mod_text, 'equal')]

                temp_results.append({
                    'segments': segments,
                    'is_heading': mod_meta.get('is_heading', False),
                    'align_type': 'match'
                })

            elif align_type == 'insert':
                mod_text = mod_texts[mod_idx]
                mod_meta = mod_paras[mod_idx][1]
                if mod_text.strip():
                    temp_results.append({
                        'segments': [(mod_text, 'insert')],
                        'is_heading': mod_meta.get('is_heading', False),
                        'align_type': 'insert',
                        'text': mod_text
                    })

            elif align_type == 'delete':
                orig_text = orig_texts[orig_idx]
                orig_meta = orig_paras[orig_idx][1]
                if orig_text.strip():
                    temp_results.append({
                        'segments': [(orig_text, 'delete')],
                        'is_heading': orig_meta.get('is_heading', False),
                        'align_type': 'delete',
                        'text': orig_text
                    })

        # Second pass: detect paragraph-level moves
        deletions = [(i, r['text']) for i, r in enumerate(temp_results)
                     if r.get('align_type') == 'delete' and len(r.get('text', '').split()) >= MIN_MOVE_WORDS]
        insertions = [(i, r['text']) for i, r in enumerate(temp_results)
                      if r.get('align_type') == 'insert' and len(r.get('text', '').split()) >= MIN_MOVE_WORDS]

        para_moves = {}  # del_idx -> ins_idx
        used_insertions = set()

        for del_idx, del_text in sorted(deletions, key=lambda x: len(x[1].split()), reverse=True):
            del_norm = normalize_text_for_move(del_text)
            best_match = None
            best_sim = 0

            for ins_idx, ins_text in insertions:
                if ins_idx in used_insertions:
                    continue
                ins_norm = normalize_text_for_move(ins_text)
                sim = calculate_similarity(del_norm, ins_norm)
                if sim >= MOVE_SIMILARITY_THRESHOLD and sim > best_sim:
                    best_sim = sim
                    best_match = ins_idx

            if best_match is not None:
                para_moves[del_idx] = best_match
                used_insertions.add(best_match)

        # Apply paragraph-level moves and build final results
        for i, result in enumerate(temp_results):
            if i in para_moves:
                # This deletion is a move source
                result['segments'] = [(result['text'], 'move_source')]
            elif i in used_insertions:
                # This insertion is a move destination
                result['segments'] = [(result['text'], 'move_dest')]

            # Count stats
            for text, seg_type in result['segments']:
                words = len(text.split())
                if seg_type == 'insert':
                    stats['insertions'] += words
                elif seg_type == 'delete':
                    stats['deletions'] += words
                elif seg_type in ('move_source', 'move_dest'):
                    stats['moves'] += words
                else:
                    stats['unchanged'] += words

            diff_results.append({
                'segments': result['segments'],
                'is_heading': result.get('is_heading', False)
            })

        # Compare tables
        if orig_tables or mod_tables:
            print("Comparing tables...")
            max_tables = max(len(orig_tables), len(mod_tables))

            for t_idx in range(max_tables):
                if t_idx < len(orig_tables) and t_idx < len(mod_tables):
                    # Both tables exist - compare them
                    print(f"  Comparing table {t_idx + 1}...")
                    table_results = diff_table(orig_tables[t_idx], mod_tables[t_idx], stats)
                    diff_results.extend(table_results)
                elif t_idx < len(mod_tables):
                    # New table - mark all as inserted
                    print(f"  Table {t_idx + 1} is new...")
                    for row in mod_tables[t_idx]:
                        row_text = ' | '.join(row)
                        diff_results.append({
                            'segments': [(row_text, 'insert')],
                            'is_heading': False,
                            'is_table_row': True
                        })
                        stats['insertions'] += len(row_text.split())
                elif t_idx < len(orig_tables):
                    # Deleted table - mark all as deleted
                    print(f"  Table {t_idx + 1} was deleted...")
                    for row in orig_tables[t_idx]:
                        row_text = ' | '.join(row)
                        diff_results.append({
                            'segments': [(row_text, 'delete')],
                            'is_heading': False,
                            'is_table_row': True
                        })
                        stats['deletions'] += len(row_text.split())

        # Generate output
        print(f"Generating {output_format} output...")

        if output_format == 'pdf':
            generate_pdf_redline(diff_results, output_path)
        else:  # word
            generate_word_redline(diff_results, output_path, modified_path)

        print(f"\nOutput saved to: {output_path}")

        return ComparisonResult(
            output_path=output_path,
            insertions=stats['insertions'],
            deletions=stats['deletions'],
            moves=stats['moves'],
            unchanged=stats['unchanged'],
            success=True
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return ComparisonResult(
            output_path=output_path,
            insertions=0,
            deletions=0,
            moves=0,
            unchanged=0,
            success=False,
            error=str(e)
        )


def generate_pdf_redline(diff_results: List[dict], output_path: str):
    """Generate a redlined PDF document."""
    generator = PdfGenerator(output_path)
    generator.generate_redline(diff_results)


def generate_word_redline(diff_results: List[dict], output_path: str, base_doc_path: Optional[str] = None):
    """Generate a redlined Word document."""
    doc = Document()

    for para_info in diff_results:
        segments = para_info.get('segments', [])
        is_heading = para_info.get('is_heading', False)

        if not segments:
            continue

        if is_heading:
            para = doc.add_heading('', level=1)
        else:
            para = doc.add_paragraph()

        for text, seg_type in segments:
            if not text:
                continue

            run = para.add_run(text)

            if seg_type == 'delete':
                run.font.strike = True
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            elif seg_type == 'insert':
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
            elif seg_type == 'move_source':
                run.font.strike = True
                run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            elif seg_type == 'move_dest':
                run.font.color.rgb = RGBColor(0, 128, 0)  # Green

    doc.save(output_path)


def main():
    """Command-line interface."""
    import argparse

    parser = argparse.ArgumentParser(
        description='Compare two documents and generate a redlined output.',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s original.docx modified.docx redline.docx
  %(prog)s original.pdf modified.pdf redline.pdf
  %(prog)s original.docx modified.pdf redline.pdf
        """
    )
    parser.add_argument('original', help='Path to original document (Word or PDF)')
    parser.add_argument('modified', help='Path to modified document (Word or PDF)')
    parser.add_argument('output', help='Path for output redlined document')
    parser.add_argument('--format', choices=['word', 'pdf'],
                       help='Output format (default: auto-detect from output extension)')

    args = parser.parse_args()

    print("=" * 60)
    print("Document Comparison Tool")
    print("=" * 60)
    print()

    result = compare_documents(
        args.original,
        args.modified,
        args.output,
        args.format
    )

    print()
    print("=" * 60)
    if result.success:
        print("COMPARISON COMPLETE")
        print("=" * 60)
        print(f"Output: {result.output_path}")
        print(f"Insertions (blue bold): {result.insertions} words")
        print(f"Deletions (red strikethrough): {result.deletions} words")
        print(f"Moves (green): {result.moves} words")
        print(f"Unchanged: {result.unchanged} words")

        total = result.insertions + result.deletions + result.moves + result.unchanged
        if total > 0:
            change_pct = (result.insertions + result.deletions + result.moves) * 100 / total
            print(f"Change percentage: {change_pct:.1f}%")
    else:
        print("COMPARISON FAILED")
        print("=" * 60)
        print(f"Error: {result.error}")
        sys.exit(1)


if __name__ == '__main__':
    main()
