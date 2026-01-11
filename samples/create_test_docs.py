"""
Create sample Word documents for testing the comparison library.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_numbered_paragraph(doc, text, level=0):
    """Add a numbered paragraph to the document."""
    paragraph = doc.add_paragraph(text)
    # Note: python-docx doesn't have full numbering support,
    # so we'll use manual numbering for the demo
    return paragraph

def create_original_document():
    """Create the original version of the contract."""
    doc = Document()

    # Title
    title = doc.add_heading('SAMPLE AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Introduction
    doc.add_paragraph('This Agreement is entered into as of January 1, 2024, by and between:')
    doc.add_paragraph('')

    # Parties
    doc.add_paragraph('Party A: Acme Corporation ("Seller")')
    doc.add_paragraph('Party B: XYZ Industries ("Buyer")')
    doc.add_paragraph('')

    # Section 1
    doc.add_heading('1. DEFINITIONS', level=1)
    doc.add_paragraph('1.1 "Agreement" means this document and all exhibits attached hereto.')
    doc.add_paragraph('1.2 "Products" means the goods described in Exhibit A.')
    doc.add_paragraph('1.3 "Services" means the services described in Exhibit B.')
    doc.add_paragraph('')

    # Section 2
    doc.add_heading('2. PURCHASE AND SALE', level=1)
    doc.add_paragraph('2.1 Subject to the terms and conditions of this Agreement, Seller agrees to sell and Buyer agrees to purchase the Products.')
    doc.add_paragraph('2.2 The purchase price shall be $50,000 USD, payable within 30 days of delivery.')
    doc.add_paragraph('2.3 Delivery shall be made to Buyer\'s facility located at 123 Main Street.')
    doc.add_paragraph('')

    # Section 3
    doc.add_heading('3. WARRANTIES', level=1)
    doc.add_paragraph('3.1 Seller warrants that the Products shall be free from defects in materials and workmanship for a period of one year from the date of delivery.')
    doc.add_paragraph('3.2 Seller warrants that the Products shall conform to the specifications set forth in Exhibit A.')
    doc.add_paragraph('')

    # Section 4
    doc.add_heading('4. LIMITATION OF LIABILITY', level=1)
    doc.add_paragraph('4.1 In no event shall either party be liable for any indirect, incidental, or consequential damages.')
    doc.add_paragraph('4.2 The total liability of Seller shall not exceed the purchase price paid by Buyer.')
    doc.add_paragraph('')

    # Section 5
    doc.add_heading('5. TERM AND TERMINATION', level=1)
    doc.add_paragraph('5.1 This Agreement shall commence on the Effective Date and continue for a period of one year.')
    doc.add_paragraph('5.2 Either party may terminate this Agreement upon 30 days written notice.')
    doc.add_paragraph('')

    # Signature block
    doc.add_paragraph('')
    doc.add_paragraph('IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.')
    doc.add_paragraph('')
    doc.add_paragraph('SELLER: _______________________')
    doc.add_paragraph('BUYER: _______________________')

    return doc

def create_modified_document():
    """Create the modified version of the contract with changes."""
    doc = Document()

    # Title (unchanged)
    title = doc.add_heading('SAMPLE AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Introduction (date changed)
    doc.add_paragraph('This Agreement is entered into as of February 15, 2024, by and between:')
    doc.add_paragraph('')

    # Parties (unchanged)
    doc.add_paragraph('Party A: Acme Corporation ("Seller")')
    doc.add_paragraph('Party B: XYZ Industries ("Buyer")')
    doc.add_paragraph('')

    # Section 1 (added a definition)
    doc.add_heading('1. DEFINITIONS', level=1)
    doc.add_paragraph('1.1 "Agreement" means this document and all exhibits attached hereto.')
    doc.add_paragraph('1.2 "Products" means the goods described in Exhibit A.')
    doc.add_paragraph('1.3 "Services" means the services described in Exhibit B.')
    doc.add_paragraph('1.4 "Confidential Information" means any proprietary information disclosed by either party.')  # NEW
    doc.add_paragraph('')

    # Section 2 (price changed, address changed)
    doc.add_heading('2. PURCHASE AND SALE', level=1)
    doc.add_paragraph('2.1 Subject to the terms and conditions of this Agreement, Seller agrees to sell and Buyer agrees to purchase the Products.')
    doc.add_paragraph('2.2 The purchase price shall be $75,000 USD, payable within 45 days of delivery.')  # CHANGED: price and payment terms
    doc.add_paragraph('2.3 Delivery shall be made to Buyer\'s facility located at 456 Commerce Drive.')  # CHANGED: address
    doc.add_paragraph('')

    # Section 3 (warranty period extended)
    doc.add_heading('3. WARRANTIES', level=1)
    doc.add_paragraph('3.1 Seller warrants that the Products shall be free from defects in materials and workmanship for a period of two years from the date of delivery.')  # CHANGED: one year -> two years
    doc.add_paragraph('3.2 Seller warrants that the Products shall conform to the specifications set forth in Exhibit A.')
    doc.add_paragraph('3.3 Seller further warrants that it has the right to sell the Products and that they are free of liens.')  # NEW
    doc.add_paragraph('')

    # Section 4 (unchanged)
    doc.add_heading('4. LIMITATION OF LIABILITY', level=1)
    doc.add_paragraph('4.1 In no event shall either party be liable for any indirect, incidental, or consequential damages.')
    doc.add_paragraph('4.2 The total liability of Seller shall not exceed the purchase price paid by Buyer.')
    doc.add_paragraph('')

    # Section 5 (termination notice changed)
    doc.add_heading('5. TERM AND TERMINATION', level=1)
    doc.add_paragraph('5.1 This Agreement shall commence on the Effective Date and continue for a period of two years.')  # CHANGED: one year -> two years
    doc.add_paragraph('5.2 Either party may terminate this Agreement upon 60 days written notice.')  # CHANGED: 30 -> 60 days
    doc.add_paragraph('5.3 Upon termination, Buyer shall return all Confidential Information to Seller.')  # NEW
    doc.add_paragraph('')

    # NEW Section 6
    doc.add_heading('6. CONFIDENTIALITY', level=1)  # NEW SECTION
    doc.add_paragraph('6.1 Each party agrees to maintain the confidentiality of all Confidential Information received from the other party.')
    doc.add_paragraph('6.2 This obligation shall survive termination of this Agreement for a period of five years.')
    doc.add_paragraph('')

    # Signature block (unchanged)
    doc.add_paragraph('')
    doc.add_paragraph('IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.')
    doc.add_paragraph('')
    doc.add_paragraph('SELLER: _______________________')
    doc.add_paragraph('BUYER: _______________________')

    return doc

if __name__ == '__main__':
    import os

    # Create samples directory path
    samples_dir = os.path.dirname(os.path.abspath(__file__))

    # Create original document
    original = create_original_document()
    original_path = os.path.join(samples_dir, 'contract_v1.docx')
    original.save(original_path)
    print(f'Created: {original_path}')

    # Create modified document
    modified = create_modified_document()
    modified_path = os.path.join(samples_dir, 'contract_v2.docx')
    modified.save(modified_path)
    print(f'Created: {modified_path}')

    print('\nSample documents created successfully!')
    print('\nKey differences between v1 and v2:')
    print('- Date changed: January 1 -> February 15')
    print('- Added definition 1.4 (Confidential Information)')
    print('- Price changed: $50,000 -> $75,000')
    print('- Payment terms: 30 days -> 45 days')
    print('- Address changed: 123 Main Street -> 456 Commerce Drive')
    print('- Warranty period: one year -> two years')
    print('- Added warranty 3.3')
    print('- Term changed: one year -> two years')
    print('- Termination notice: 30 days -> 60 days')
    print('- Added termination clause 5.3')
    print('- Added new Section 6 (Confidentiality)')
