"""
Document comparison that PRESERVES original formatting.

Strategy:
1. Copy the MODIFIED document as the base (keeps all formatting intact)
2. For each paragraph, find what was deleted and insert it with strikethrough
3. Mark inserted text with blue bold
4. Detect moved text and show in green (source struck through, destination plain green)
"""
import difflib
import re
import copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import os
import shutil
from typing import List, Tuple, Dict, Optional
from dataclasses import dataclass

# Word XML namespaces
WORD_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

# Move detection threshold - text must be this similar to be considered a move
MOVE_SIMILARITY_THRESHOLD = 0.85

# Minimum word count for move detection (very short phrases are likely coincidental)
MIN_MOVE_WORDS = 3


@dataclass
class MoveCandidate:
    """Represents a potential moved text block."""
    text: str
    paragraph_idx: int
    is_deletion: bool  # True if from original (deletion), False if from modified (insertion)
    word_count: int


@dataclass
class TableCell:
    """Represents a table cell for comparison."""
    text: str
    row_idx: int
    col_idx: int
    row_span: int = 1
    col_span: int = 1


@dataclass
class TableInfo:
    """Represents a table extracted from a document."""
    table_idx: int
    rows: List[List[TableCell]]
    row_count: int
    col_count: int

def get_paragraph_text(para):
    """Get plain text from a paragraph."""
    return para.text or ''

def tokenize(text):
    """Split text into words while preserving whitespace."""
    words = []
    for match in re.finditer(r'\S+|\s+', text):
        words.append(match.group())
    return words

def diff_texts(original_text, modified_text):
    """
    Compute word-level diff between two texts.
    Returns list of (text, type) where type is 'equal', 'delete', or 'insert'
    """
    orig_words = tokenize(original_text)
    mod_words = tokenize(modified_text)

    matcher = difflib.SequenceMatcher(None, orig_words, mod_words, autojunk=False)
    result = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            text = ''.join(orig_words[i1:i2])
            if text:
                result.append((text, 'equal'))
        elif tag == 'delete':
            text = ''.join(orig_words[i1:i2])
            if text:
                result.append((text, 'delete'))
        elif tag == 'insert':
            text = ''.join(mod_words[j1:j2])
            if text:
                result.append((text, 'insert'))
        elif tag == 'replace':
            del_text = ''.join(orig_words[i1:i2])
            ins_text = ''.join(mod_words[j1:j2])
            if del_text:
                result.append((del_text, 'delete'))
            if ins_text:
                result.append((ins_text, 'insert'))

    return result

def get_first_run_formatting(para):
    """Extract formatting from the first run of a paragraph."""
    formatting = {}
    if para.runs:
        run = para.runs[0]
        formatting['bold'] = run.bold
        formatting['italic'] = run.italic
        formatting['underline'] = run.underline
        formatting['font_name'] = run.font.name
        formatting['font_size'] = run.font.size
    return formatting

def apply_base_formatting(run, formatting):
    """Apply base formatting to a run."""
    if formatting.get('font_name'):
        run.font.name = formatting['font_name']
    if formatting.get('font_size'):
        run.font.size = formatting['font_size']
    # Don't copy bold/italic as we may override them

def clear_paragraph_content(para):
    """Clear all run content from paragraph, preserving paragraph properties."""
    p_element = para._p

    # Find and remove all content elements (runs, hyperlinks, bookmarks, etc.)
    # Keep only pPr (paragraph properties)
    elements_to_remove = []
    for child in p_element:
        tag = child.tag
        # Keep paragraph properties (pPr), remove everything else
        if not tag.endswith('}pPr'):
            elements_to_remove.append(child)

    for elem in elements_to_remove:
        p_element.remove(elem)

def rebuild_paragraph_with_diff(para, diff_segments, base_formatting=None):
    """
    Completely rebuild paragraph content with diff segments.
    """
    # First, completely clear the paragraph of all runs
    clear_paragraph_content(para)

    # Now add new runs for each segment
    for text, seg_type in diff_segments:
        if not text:
            continue

        run = para.add_run(text)

        # Apply base formatting
        if base_formatting:
            apply_base_formatting(run, base_formatting)

        # Apply diff formatting
        if seg_type == 'delete':
            run.font.strike = True
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        elif seg_type == 'insert':
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
        elif seg_type == 'move_source':
            # Moved text at original location - green strikethrough
            run.font.strike = True
            run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif seg_type == 'move_dest':
            # Moved text at new location - green (no strikethrough)
            run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        # 'equal' keeps normal formatting

def apply_formatting_to_existing_runs(para, formatting_type):
    """Apply redline formatting to all existing runs in a paragraph."""
    for run in para.runs:
        if formatting_type == 'delete':
            run.font.strike = True
            run.font.color.rgb = RGBColor(255, 0, 0)
        elif formatting_type == 'insert':
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 255)
        elif formatting_type == 'move_source':
            run.font.strike = True
            run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif formatting_type == 'move_dest':
            run.font.color.rgb = RGBColor(0, 128, 0)  # Green

def calculate_similarity(text1, text2):
    """Calculate similarity between two texts using multiple methods."""
    if text1 == text2:
        return 1.0
    if not text1 or not text2:
        return 0.0

    text1 = text1.strip()
    text2 = text2.strip()

    if not text1 and not text2:
        return 1.0
    if not text1 or not text2:
        return 0.0

    # Method 1: Word-based Jaccard similarity
    words1 = set(text1.lower().split())
    words2 = set(text2.lower().split())

    word_sim = 0.0
    if words1 and words2:
        intersection = len(words1 & words2)
        union = len(words1 | words2)
        word_sim = intersection / union if union > 0 else 0.0

    # Method 2: Character sequence similarity (better for small edits)
    seq_sim = difflib.SequenceMatcher(None, text1.lower(), text2.lower()).ratio()

    # Use the higher of the two similarities
    return max(word_sim, seq_sim)


def normalize_text_for_move(text: str) -> str:
    """Normalize text for move detection comparison."""
    # Lowercase, collapse whitespace, strip
    text = text.lower().strip()
    text = re.sub(r'\s+', ' ', text)
    return text


def detect_moves(deletions: List[MoveCandidate], insertions: List[MoveCandidate]) -> Dict[int, int]:
    """
    Detect moved text by matching deletions with insertions.

    Returns a dict mapping deletion paragraph index to insertion paragraph index
    for detected moves.
    """
    moves = {}  # deletion_idx -> insertion_idx
    used_insertions = set()

    # Sort by word count descending - match longer phrases first (more confident)
    sorted_deletions = sorted(
        [(i, d) for i, d in enumerate(deletions) if d.word_count >= MIN_MOVE_WORDS],
        key=lambda x: x[1].word_count,
        reverse=True
    )

    for del_list_idx, deletion in sorted_deletions:
        del_normalized = normalize_text_for_move(deletion.text)
        best_match = None
        best_similarity = 0

        for ins_list_idx, insertion in enumerate(insertions):
            if ins_list_idx in used_insertions:
                continue
            if insertion.word_count < MIN_MOVE_WORDS:
                continue

            ins_normalized = normalize_text_for_move(insertion.text)

            # Calculate similarity
            similarity = calculate_similarity(del_normalized, ins_normalized)

            if similarity >= MOVE_SIMILARITY_THRESHOLD and similarity > best_similarity:
                best_similarity = similarity
                best_match = ins_list_idx

        if best_match is not None:
            moves[deletion.paragraph_idx] = insertions[best_match].paragraph_idx
            used_insertions.add(best_match)

    return moves


def detect_word_level_moves(diff_segments: List[Tuple[str, str]]) -> List[Tuple[str, str]]:
    """
    Detect moves within word-level diff segments.

    Takes diff segments and identifies deletions that match insertions,
    converting them to move_source/move_dest types.
    """
    # Collect deletions and insertions
    deletions = []  # (index, text, normalized)
    insertions = []  # (index, text, normalized)

    for i, (text, seg_type) in enumerate(diff_segments):
        words = len(text.split())
        if seg_type == 'delete' and words >= MIN_MOVE_WORDS:
            deletions.append((i, text, normalize_text_for_move(text)))
        elif seg_type == 'insert' and words >= MIN_MOVE_WORDS:
            insertions.append((i, text, normalize_text_for_move(text)))

    if not deletions or not insertions:
        return diff_segments

    # Find matching moves
    moves = {}  # deletion_idx -> insertion_idx
    used_insertions = set()

    # Sort deletions by word count descending
    sorted_deletions = sorted(deletions, key=lambda x: len(x[1].split()), reverse=True)

    for del_idx, del_text, del_norm in sorted_deletions:
        best_match = None
        best_similarity = 0

        for ins_idx, ins_text, ins_norm in insertions:
            if ins_idx in used_insertions:
                continue

            similarity = calculate_similarity(del_norm, ins_norm)

            if similarity >= MOVE_SIMILARITY_THRESHOLD and similarity > best_similarity:
                best_similarity = similarity
                best_match = ins_idx

        if best_match is not None:
            moves[del_idx] = best_match
            used_insertions.add(best_match)

    if not moves:
        return diff_segments

    # Create new segments with move markers
    result = []
    for i, (text, seg_type) in enumerate(diff_segments):
        if i in moves:
            # This deletion is a move source
            result.append((text, 'move_source'))
        elif i in used_insertions:
            # This insertion is a move destination
            result.append((text, 'move_dest'))
        else:
            result.append((text, seg_type))

    return result


def align_paragraphs(orig_paras, mod_paras):
    """Align paragraphs between documents using LCS."""
    m, n = len(orig_paras), len(mod_paras)

    # Build LCS table
    lcs = [[0] * (n + 1) for _ in range(m + 1)]

    for i in range(1, m + 1):
        for j in range(1, n + 1):
            orig_text = get_paragraph_text(orig_paras[i-1])
            mod_text = get_paragraph_text(mod_paras[j-1])

            if calculate_similarity(orig_text, mod_text) >= 0.4:
                lcs[i][j] = lcs[i-1][j-1] + 1
            else:
                lcs[i][j] = max(lcs[i-1][j], lcs[i][j-1])

    # Backtrack
    alignments = []
    i, j = m, n

    while i > 0 or j > 0:
        if i > 0 and j > 0:
            orig_text = get_paragraph_text(orig_paras[i-1])
            mod_text = get_paragraph_text(mod_paras[j-1])
            if calculate_similarity(orig_text, mod_text) >= 0.4:
                alignments.append((i-1, j-1, 'match'))
                i -= 1
                j -= 1
                continue

        if j > 0 and (i == 0 or lcs[i][j-1] >= lcs[i-1][j]):
            alignments.append((-1, j-1, 'insert'))
            j -= 1
        else:
            alignments.append((i-1, -1, 'delete'))
            i -= 1

    alignments.reverse()
    return alignments


# =============================================================================
# TABLE COMPARISON FUNCTIONS
# =============================================================================

def extract_table_info(table, table_idx: int) -> TableInfo:
    """Extract table structure and content from a Word table."""
    rows = []
    for row_idx, row in enumerate(table.rows):
        row_cells = []
        for col_idx, cell in enumerate(row.cells):
            # Get cell text (join all paragraphs)
            cell_text = '\n'.join(p.text for p in cell.paragraphs)
            row_cells.append(TableCell(
                text=cell_text,
                row_idx=row_idx,
                col_idx=col_idx
            ))
        rows.append(row_cells)

    row_count = len(table.rows)
    col_count = len(table.rows[0].cells) if table.rows else 0

    return TableInfo(
        table_idx=table_idx,
        rows=rows,
        row_count=row_count,
        col_count=col_count
    )


def get_row_text(row_cells: List[TableCell]) -> str:
    """Get combined text from a row for comparison."""
    return ' | '.join(cell.text for cell in row_cells)


def align_table_rows(orig_rows: List[List[TableCell]], mod_rows: List[List[TableCell]]) -> List[Tuple[int, int, str]]:
    """Align rows between two tables using LCS algorithm."""
    m, n = len(orig_rows), len(mod_rows)

    # Build LCS table
    lcs = [[0] * (n + 1) for _ in range(m + 1)]

    for i in range(1, m + 1):
        for j in range(1, n + 1):
            orig_text = get_row_text(orig_rows[i-1])
            mod_text = get_row_text(mod_rows[j-1])

            if calculate_similarity(orig_text, mod_text) >= 0.4:
                lcs[i][j] = lcs[i-1][j-1] + 1
            else:
                lcs[i][j] = max(lcs[i-1][j], lcs[i][j-1])

    # Backtrack
    alignments = []
    i, j = m, n

    while i > 0 or j > 0:
        if i > 0 and j > 0:
            orig_text = get_row_text(orig_rows[i-1])
            mod_text = get_row_text(mod_rows[j-1])
            if calculate_similarity(orig_text, mod_text) >= 0.4:
                alignments.append((i-1, j-1, 'match'))
                i -= 1
                j -= 1
                continue

        if j > 0 and (i == 0 or lcs[i][j-1] >= lcs[i-1][j]):
            alignments.append((-1, j-1, 'insert'))
            j -= 1
        else:
            alignments.append((i-1, -1, 'delete'))
            i -= 1

    alignments.reverse()
    return alignments


def compare_table_cells(orig_table, out_table, stats):
    """
    Compare two tables cell by cell and apply diff formatting.

    Args:
        orig_table: Original table (from original document)
        out_table: Output table (from modified document, will be modified in place)
        stats: Dictionary to track change statistics
    """
    orig_info = extract_table_info(orig_table, 0)
    out_info = extract_table_info(out_table, 0)

    # Align rows
    row_alignments = align_table_rows(orig_info.rows, out_info.rows)

    for orig_row_idx, mod_row_idx, align_type in row_alignments:
        if align_type == 'match' and orig_row_idx >= 0 and mod_row_idx >= 0:
            orig_row = orig_info.rows[orig_row_idx]
            out_row = out_table.rows[mod_row_idx]

            # Compare each cell in the row
            max_cols = max(len(orig_row), len(out_row.cells))

            for col_idx in range(max_cols):
                if col_idx < len(orig_row) and col_idx < len(out_row.cells):
                    # Both cells exist - compare them
                    orig_cell_text = orig_row[col_idx].text
                    out_cell = out_row.cells[col_idx]
                    out_cell_text = '\n'.join(p.text for p in out_cell.paragraphs)

                    if orig_cell_text.strip() != out_cell_text.strip():
                        # Cell content differs - apply diff
                        compare_cell_content(orig_cell_text, out_cell, stats)
                    else:
                        stats['unchanged'] += len(out_cell_text.split())

                elif col_idx < len(out_row.cells):
                    # New column in modified - mark as inserted
                    out_cell = out_row.cells[col_idx]
                    mark_cell_as_inserted(out_cell, stats)

        elif align_type == 'insert' and mod_row_idx >= 0:
            # Entire row is new - mark all cells as inserted
            out_row = out_table.rows[mod_row_idx]
            for cell in out_row.cells:
                mark_cell_as_inserted(cell, stats)

        elif align_type == 'delete' and orig_row_idx >= 0:
            # Row was deleted - we can't easily show this since we're using modified as base
            # Count the deleted words
            orig_row = orig_info.rows[orig_row_idx]
            for cell in orig_row:
                stats['deletions'] += len(cell.text.split())


def compare_cell_content(orig_text: str, out_cell, stats):
    """Compare cell content and apply diff formatting to the cell."""
    # Get the cell's paragraphs
    for para_idx, para in enumerate(out_cell.paragraphs):
        # For simplicity, compare paragraph by paragraph
        # Get corresponding original text (split by newlines)
        orig_lines = orig_text.split('\n')
        mod_text = para.text

        if para_idx < len(orig_lines):
            orig_para_text = orig_lines[para_idx]
        else:
            orig_para_text = ''

        if orig_para_text.strip() != mod_text.strip():
            # Get base formatting
            base_formatting = get_first_run_formatting(para)

            # Compute diff
            diff_segments = diff_texts(orig_para_text, mod_text)

            # Apply move detection
            diff_segments = detect_word_level_moves(diff_segments)

            # Count changes
            for text, seg_type in diff_segments:
                words = len(text.split())
                if seg_type == 'insert':
                    stats['insertions'] += words
                elif seg_type == 'delete':
                    stats['deletions'] += words
                elif seg_type in ('move_source', 'move_dest'):
                    stats['moves'] += words
                else:
                    stats['unchanged'] += words

            # Rebuild paragraph with diff
            rebuild_paragraph_with_diff(para, diff_segments, base_formatting)
        else:
            stats['unchanged'] += len(mod_text.split())


def mark_cell_as_inserted(cell, stats):
    """Mark all content in a cell as inserted (blue bold)."""
    for para in cell.paragraphs:
        text = para.text
        if text.strip():
            apply_formatting_to_existing_runs(para, 'insert')
            stats['insertions'] += len(text.split())


def mark_cell_as_deleted(cell, stats):
    """Mark all content in a cell as deleted (red strikethrough)."""
    for para in cell.paragraphs:
        text = para.text
        if text.strip():
            apply_formatting_to_existing_runs(para, 'delete')
            stats['deletions'] += len(text.split())


def compare_tables_list(orig_tables, out_tables, stats):
    """
    Compare lists of tables between original and output documents.

    Uses position-based matching (table 1 with table 1, etc.)
    """
    max_tables = max(len(orig_tables), len(out_tables))

    for i in range(max_tables):
        if i < len(orig_tables) and i < len(out_tables):
            # Both tables exist - compare them
            print(f"  Comparing table {i + 1}...")
            compare_table_cells(orig_tables[i], out_tables[i], stats)
        elif i < len(out_tables):
            # New table in modified - mark all cells as inserted
            print(f"  Table {i + 1} is new (marking as inserted)...")
            out_table = out_tables[i]
            for row in out_table.rows:
                for cell in row.cells:
                    mark_cell_as_inserted(cell, stats)


def compare_paragraphs_list(orig_paras, out_paras, stats, context="", detect_moves_flag=True):
    """
    Compare two lists of paragraphs and apply diff formatting.
    Includes move detection at both paragraph and word level.
    """
    if not orig_paras and not out_paras:
        return

    # Align paragraphs
    alignments = align_paragraphs(orig_paras, out_paras)

    # First pass: collect deletions and insertions for paragraph-level move detection
    if detect_moves_flag:
        deletion_candidates = []
        insertion_candidates = []

        for orig_idx, mod_idx, align_type in alignments:
            if align_type == 'delete' and orig_idx >= 0:
                text = get_paragraph_text(orig_paras[orig_idx])
                if text.strip():
                    deletion_candidates.append(MoveCandidate(
                        text=text,
                        paragraph_idx=orig_idx,
                        is_deletion=True,
                        word_count=len(text.split())
                    ))
            elif align_type == 'insert' and mod_idx >= 0:
                text = get_paragraph_text(out_paras[mod_idx])
                if text.strip():
                    insertion_candidates.append(MoveCandidate(
                        text=text,
                        paragraph_idx=mod_idx,
                        is_deletion=False,
                        word_count=len(text.split())
                    ))

        # Detect paragraph-level moves
        para_moves = detect_moves(deletion_candidates, insertion_candidates)
        # para_moves: orig_idx -> mod_idx

        # Build reverse lookup: mod_idx -> orig_idx (for move destinations)
        move_destinations = {v: k for k, v in para_moves.items()}
    else:
        para_moves = {}
        move_destinations = {}

    # Second pass: apply formatting
    for orig_idx, mod_idx, align_type in alignments:
        if align_type == 'match' and mod_idx >= 0 and orig_idx >= 0:
            orig_para = orig_paras[orig_idx]
            out_para = out_paras[mod_idx]

            orig_text = get_paragraph_text(orig_para)
            mod_text = get_paragraph_text(out_para)

            # Only process if there's actual text and it differs
            if orig_text.strip() != mod_text.strip() and (orig_text.strip() or mod_text.strip()):
                # Get base formatting from output paragraph
                base_formatting = get_first_run_formatting(out_para)

                # Compute diff
                diff_segments = diff_texts(orig_text, mod_text)

                # Apply word-level move detection
                if detect_moves_flag:
                    diff_segments = detect_word_level_moves(diff_segments)

                # Count changes
                for text, seg_type in diff_segments:
                    words = len(text.split())
                    if seg_type == 'insert':
                        stats['insertions'] += words
                    elif seg_type == 'delete':
                        stats['deletions'] += words
                    elif seg_type in ('move_source', 'move_dest'):
                        stats['moves'] += words
                    else:
                        stats['unchanged'] += words

                # Rebuild paragraph with diff formatting
                rebuild_paragraph_with_diff(out_para, diff_segments, base_formatting)
            else:
                stats['unchanged'] += len(mod_text.split())

        elif align_type == 'insert' and mod_idx >= 0:
            out_para = out_paras[mod_idx]
            text = get_paragraph_text(out_para)
            if text.strip():
                # Check if this is a move destination
                if mod_idx in move_destinations:
                    apply_formatting_to_existing_runs(out_para, 'move_dest')
                    stats['moves'] += len(text.split())
                else:
                    apply_formatting_to_existing_runs(out_para, 'insert')
                    stats['insertions'] += len(text.split())

        elif align_type == 'delete' and orig_idx >= 0:
            orig_para = orig_paras[orig_idx]
            text = get_paragraph_text(orig_para)
            if text.strip():
                # Check if this is a move source
                if orig_idx in para_moves:
                    stats['moves'] += len(text.split())
                    # Note: We can't easily add the source to output doc since we're using
                    # modified doc as base. The move will show at destination in green.
                else:
                    stats['deletions'] += len(text.split())


def compare_with_full_formatting(original_path, modified_path, output_path):
    """
    Compare documents while preserving formatting.
    Uses modified document as base and applies diff formatting.
    Compares body, headers, and footers.
    """
    print("Loading documents...")

    # Start by copying the modified file as our base
    shutil.copy2(modified_path, output_path)

    # Now open both original and output
    original_doc = Document(original_path)
    output_doc = Document(output_path)

    stats = {'insertions': 0, 'deletions': 0, 'unchanged': 0, 'moves': 0}

    # Compare main body paragraphs
    orig_paras = list(original_doc.paragraphs)
    out_paras = list(output_doc.paragraphs)

    print(f"Original: {len(orig_paras)} body paragraphs")
    print(f"Modified: {len(out_paras)} body paragraphs")

    print("Comparing body paragraphs...")
    compare_paragraphs_list(orig_paras, out_paras, stats, "body")

    # Compare tables
    orig_tables = list(original_doc.tables)
    out_tables = list(output_doc.tables)

    if orig_tables or out_tables:
        print(f"Tables: {len(orig_tables)} original, {len(out_tables)} modified")
        print("Comparing tables...")
        compare_tables_list(orig_tables, out_tables, stats)

    # Compare headers and footers for each section
    orig_sections = list(original_doc.sections)
    out_sections = list(output_doc.sections)

    print(f"Sections: {len(orig_sections)} original, {len(out_sections)} modified")

    # Process headers and footers
    for i in range(min(len(orig_sections), len(out_sections))):
        orig_section = orig_sections[i]
        out_section = out_sections[i]

        # Compare headers
        try:
            if orig_section.header and out_section.header:
                orig_header_paras = list(orig_section.header.paragraphs)
                out_header_paras = list(out_section.header.paragraphs)
                if orig_header_paras or out_header_paras:
                    print(f"Comparing header (section {i+1})...")
                    compare_paragraphs_list(orig_header_paras, out_header_paras, stats, f"header_{i}")
        except Exception as e:
            print(f"  Warning: Could not compare header section {i+1}: {e}")

        # Compare first page header if different
        try:
            if orig_section.first_page_header and out_section.first_page_header:
                orig_fp_header_paras = list(orig_section.first_page_header.paragraphs)
                out_fp_header_paras = list(out_section.first_page_header.paragraphs)
                if orig_fp_header_paras or out_fp_header_paras:
                    print(f"Comparing first page header (section {i+1})...")
                    compare_paragraphs_list(orig_fp_header_paras, out_fp_header_paras, stats, f"fp_header_{i}")
        except Exception as e:
            pass  # First page header may not exist

        # Compare footers
        try:
            if orig_section.footer and out_section.footer:
                orig_footer_paras = list(orig_section.footer.paragraphs)
                out_footer_paras = list(out_section.footer.paragraphs)
                if orig_footer_paras or out_footer_paras:
                    print(f"Comparing footer (section {i+1})...")
                    compare_paragraphs_list(orig_footer_paras, out_footer_paras, stats, f"footer_{i}")
        except Exception as e:
            print(f"  Warning: Could not compare footer section {i+1}: {e}")

        # Compare first page footer if different
        try:
            if orig_section.first_page_footer and out_section.first_page_footer:
                orig_fp_footer_paras = list(orig_section.first_page_footer.paragraphs)
                out_fp_footer_paras = list(out_section.first_page_footer.paragraphs)
                if orig_fp_footer_paras or out_fp_footer_paras:
                    print(f"Comparing first page footer (section {i+1})...")
                    compare_paragraphs_list(orig_fp_footer_paras, out_fp_footer_paras, stats, f"fp_footer_{i}")
        except Exception as e:
            pass  # First page footer may not exist

    print(f"Saving to {output_path}...")
    output_doc.save(output_path)

    return stats

if __name__ == '__main__':
    import sys

    if len(sys.argv) >= 4:
        original = sys.argv[1]
        modified = sys.argv[2]
        output = sys.argv[3]
    else:
        # Default test files
        samples_dir = os.path.dirname(os.path.abspath(__file__))
        original = os.path.join(samples_dir, 'contract_v1.docx')
        modified = os.path.join(samples_dir, 'contract_v2.docx')
        output = os.path.join(samples_dir, 'contract_redline_formatted.docx')

    print("="*60)
    print("Document Comparison (Formatting Preserved)")
    print("="*60)
    print(f"Original: {original}")
    print(f"Modified: {modified}")
    print(f"Output:   {output}")
    print()

    stats = compare_with_full_formatting(original, modified, output)

    print()
    print("="*60)
    print("COMPLETE")
    print("="*60)
    print(f"Output: {output}")
    print(f"Insertions (blue bold): {stats['insertions']} words")
    print(f"Deletions (red strikethrough): {stats['deletions']} words")
    print(f"Moves (green): {stats['moves']} words")
    print(f"Unchanged: {stats['unchanged']} words")
