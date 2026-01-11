"""Create test documents with moved paragraphs for testing move detection."""

from docx import Document
import os

def create_test_docs():
    samples_dir = os.path.dirname(os.path.abspath(__file__))

    # Original document
    doc1 = Document()
    doc1.add_heading('Test Document for Move Detection', 0)

    doc1.add_paragraph('This is the first paragraph that will remain unchanged.')

    doc1.add_paragraph('This paragraph contains important information that will be moved to the end of the document in the modified version.')

    doc1.add_paragraph('Here is some middle content that stays in place.')

    doc1.add_paragraph('Another paragraph with some regular content.')

    doc1.add_paragraph('This is the last paragraph in the original document.')

    output1 = os.path.join(samples_dir, 'move_test_original.docx')
    doc1.save(output1)
    print(f"Created: {output1}")

    # Modified document (with moved paragraph)
    doc2 = Document()
    doc2.add_heading('Test Document for Move Detection', 0)

    doc2.add_paragraph('This is the first paragraph that will remain unchanged.')

    # The important paragraph is removed from here

    doc2.add_paragraph('Here is some middle content that stays in place.')

    doc2.add_paragraph('Another paragraph with some regular content.')

    doc2.add_paragraph('This is the last paragraph in the original document.')

    # The important paragraph is added at the end (moved)
    doc2.add_paragraph('This paragraph contains important information that will be moved to the end of the document in the modified version.')

    output2 = os.path.join(samples_dir, 'move_test_modified.docx')
    doc2.save(output2)
    print(f"Created: {output2}")

    print("\nTest documents created successfully!")
    print("Run comparison with:")
    print(f"  python compare_preserve_formatting.py {output1} {output2} move_test_redline.docx")


if __name__ == '__main__':
    create_test_docs()
