"""Create test documents with tables for testing table comparison."""

from docx import Document
from docx.shared import Inches
import os

def create_test_docs():
    samples_dir = os.path.dirname(os.path.abspath(__file__))

    # Original document with a table
    doc1 = Document()
    doc1.add_heading('Test Document with Tables', 0)
    doc1.add_paragraph('This document contains a table for testing cell-level comparison.')

    # Add a 3x3 table
    table1 = doc1.add_table(rows=3, cols=3)
    table1.style = 'Table Grid'

    # Header row
    table1.cell(0, 0).text = 'Name'
    table1.cell(0, 1).text = 'Role'
    table1.cell(0, 2).text = 'Department'

    # Data rows
    table1.cell(1, 0).text = 'John Smith'
    table1.cell(1, 1).text = 'Software Engineer'
    table1.cell(1, 2).text = 'Engineering'

    table1.cell(2, 0).text = 'Jane Doe'
    table1.cell(2, 1).text = 'Product Manager'
    table1.cell(2, 2).text = 'Product'

    doc1.add_paragraph('End of document.')

    output1 = os.path.join(samples_dir, 'table_test_original.docx')
    doc1.save(output1)
    print(f"Created: {output1}")

    # Modified document with changes
    doc2 = Document()
    doc2.add_heading('Test Document with Tables', 0)
    doc2.add_paragraph('This document contains a table for testing cell-level comparison.')

    # Add a 4x3 table (added a new row)
    table2 = doc2.add_table(rows=4, cols=3)
    table2.style = 'Table Grid'

    # Header row (unchanged)
    table2.cell(0, 0).text = 'Name'
    table2.cell(0, 1).text = 'Role'
    table2.cell(0, 2).text = 'Department'

    # Data row 1 - modified role
    table2.cell(1, 0).text = 'John Smith'
    table2.cell(1, 1).text = 'Senior Software Engineer'  # Changed from "Software Engineer"
    table2.cell(1, 2).text = 'Engineering'

    # Data row 2 - modified department
    table2.cell(2, 0).text = 'Jane Doe'
    table2.cell(2, 1).text = 'Product Manager'
    table2.cell(2, 2).text = 'Product Management'  # Changed from "Product"

    # New row added
    table2.cell(3, 0).text = 'Bob Wilson'
    table2.cell(3, 1).text = 'Designer'
    table2.cell(3, 2).text = 'Design'

    doc2.add_paragraph('End of document.')

    output2 = os.path.join(samples_dir, 'table_test_modified.docx')
    doc2.save(output2)
    print(f"Created: {output2}")

    print("\nTest documents created successfully!")
    print("\nExpected changes:")
    print("  - Cell (1,1): 'Software Engineer' -> 'Senior Software Engineer' (insertion)")
    print("  - Cell (2,2): 'Product' -> 'Product Management' (insertion)")
    print("  - Row 3: New row added (Bob Wilson, Designer, Design)")
    print("\nRun comparison with:")
    print(f"  python compare_preserve_formatting.py {output1} {output2} table_test_redline.docx")


if __name__ == '__main__':
    create_test_docs()
